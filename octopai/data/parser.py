"""
Resource Parser - Enhanced Multi-Format Resource Processor

This module provides intelligent parsers for various file formats to convert them into
skill resources that can be used for skill creation and evolution.

Enhanced Features:
- Smart content extraction and structuring
- Content quality assessment
- Metadata enrichment
- Multi-format support with fallback mechanisms
- Content summarization capabilities
"""

import os
import base64
import mimetypes
from typing import Dict, Any, Optional, List, Tuple
from dataclasses import dataclass, field
from enum import Enum
from datetime import datetime


class ResourceType(Enum):
    """Types of resources that can be parsed"""
    PDF = "pdf"
    DOC = "doc"
    DOCX = "docx"
    EXCEL = "excel"
    IMAGE = "image"
    VIDEO = "video"
    TEXT = "text"
    MARKDOWN = "markdown"
    HTML = "html"
    JSON = "json"
    YAML = "yaml"
    CSV = "csv"
    XML = "xml"
    ZIP = "zip"
    UNKNOWN = "unknown"


class ContentQuality(Enum):
    """Quality assessment for parsed content"""
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    EXCELLENT = "excellent"


@dataclass
class ResourceMetadata:
    """Enhanced metadata for parsed resources"""
    file_name: str
    file_size: int
    file_type: str
    mime_type: Optional[str] = None
    created_at: Optional[str] = None
    modified_at: Optional[str] = None
    content_length: int = 0
    content_quality: ContentQuality = ContentQuality.MEDIUM
    language: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    summary: Optional[str] = None
    custom_fields: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "file_name": self.file_name,
            "file_size": self.file_size,
            "file_type": self.file_type,
            "mime_type": self.mime_type,
            "created_at": self.created_at,
            "modified_at": self.modified_at,
            "content_length": self.content_length,
            "content_quality": self.content_quality.value,
            "language": self.language,
            "keywords": self.keywords,
            "summary": self.summary,
            "custom_fields": self.custom_fields
        }


@dataclass
class ParsedResource:
    """Represents a parsed resource ready for skill usage with enhanced features"""
    file_path: str
    resource_type: ResourceType
    text_content: str
    metadata: ResourceMetadata
    images: List[bytes] = field(default_factory=list)
    raw_data: Optional[bytes] = None
    structured_sections: Dict[str, str] = field(default_factory=dict)
    tables: List[List[List[str]]] = field(default_factory=list)
    links: List[Tuple[str, str]] = field(default_factory=list)
    
    def __post_init__(self):
        if self.images is None:
            self.images = []
        if self.structured_sections is None:
            self.structured_sections = {}
        if self.tables is None:
            self.tables = []
        if self.links is None:
            self.links = []
    
    def to_skill_resource(self) -> str:
        """Convert parsed resource to skill resource format with enhanced structure"""
        parts = []
        
        parts.append(f"# Resource: {self.metadata.file_name}")
        parts.append(f"**Type**: {self.resource_type.value}")
        parts.append(f"**Quality**: {self.metadata.content_quality.value}")
        parts.append("")
        
        if self.metadata.summary:
            parts.append("## Summary")
            parts.append(self.metadata.summary)
            parts.append("")
        
        if self.metadata.keywords:
            parts.append("## Keywords")
            parts.append(", ".join(self.metadata.keywords))
            parts.append("")
        
        if self.structured_sections:
            for section_name, section_content in self.structured_sections.items():
                parts.append(f"## {section_name}")
                parts.append(section_content)
                parts.append("")
        elif self.text_content:
            parts.append("## Content")
            parts.append(self.text_content)
        
        if self.tables:
            parts.append("## Tables")
            for i, table in enumerate(self.tables, 1):
                parts.append(f"### Table {i}")
                for row in table:
                    parts.append("| " + " | ".join(row) + " |")
                parts.append("")
        
        if self.images:
            parts.append(f"## Images: {len(self.images)} image(s) extracted")
        
        parts.append("")
        parts.append("---")
        parts.append("## Technical Metadata")
        meta_dict = self.metadata.to_dict()
        for key, value in meta_dict.items():
            if value and key != "custom_fields":
                parts.append(f"- **{key.replace('_', ' ').title()}**: {value}")
        
        return "\n".join(parts)
    
    def get_base64_images(self) -> List[str]:
        """Get images as base64 encoded strings"""
        return [base64.b64encode(img).decode('utf-8') for img in self.images]
    
    def get_content_summary(self, max_length: int = 500) -> str:
        """Get a summary of the content"""
        if self.metadata.summary:
            return self.metadata.summary
        
        content = self.text_content[:max_length]
        if len(self.text_content) > max_length:
            content += "..."
        return content


class ContentAnalyzer:
    """Content analyzer for quality assessment and metadata enrichment"""
    
    def __init__(self):
        self.common_keywords = set([
            'important', 'note', 'warning', 'caution', 'tip', 'example',
            'usage', 'installation', 'configuration', 'setup', 'troubleshooting',
            'best', 'practice', 'recommendation', 'guide', 'tutorial', 'reference',
            'api', 'function', 'method', 'class', 'parameter', 'return', 'error'
        ])
    
    def analyze_content(self, content: str, file_path: str) -> Tuple[ContentQuality, List[str], Optional[str]]:
        """
        Analyze content quality and extract keywords
        
        Args:
            content: Text content to analyze
            file_path: File path for context
            
        Returns:
            Tuple of (quality_level, keywords, summary)
        """
        quality = ContentQuality.MEDIUM
        keywords = []
        summary = None
        
        content_lower = content.lower()
        content_length = len(content)
        
        if content_length < 100:
            quality = ContentQuality.LOW
        elif content_length > 5000:
            quality = ContentQuality.HIGH
        
        sections = self._extract_sections(content)
        if len(sections) >= 3:
            quality = ContentQuality.HIGH
        if len(sections) >= 5:
            quality = ContentQuality.EXCELLENT
        
        has_code = '```' in content
        has_examples = 'example' in content_lower or '## usage' in content_lower
        has_troubleshooting = 'troubleshoot' in content_lower or 'faq' in content_lower
        has_best_practices = 'best practice' in content_lower or 'recommendation' in content_lower
        
        if has_code and has_examples and has_troubleshooting:
            quality = ContentQuality.EXCELLENT
        
        keywords = self._extract_keywords(content_lower)
        summary = self._generate_summary(content, sections)
        
        return quality, keywords, summary
    
    def _extract_sections(self, content: str) -> List[str]:
        """Extract section headings from content"""
        sections = []
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('#'):
                sections.append(line.lstrip('#').strip())
        return sections
    
    def _extract_keywords(self, content_lower: str) -> List[str]:
        """Extract keywords from content"""
        words = content_lower.split()
        found_keywords = []
        for word in words:
            word = word.strip('.,;:!?"\'()[]{}')
            if word in self.common_keywords and word not in found_keywords:
                found_keywords.append(word)
        return found_keywords[:10]
    
    def _generate_summary(self, content: str, sections: List[str]) -> Optional[str]:
        """Generate a simple content summary"""
        if not content:
            return None
        
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        
        if lines:
            first_paragraph = []
            for line in lines[:5]:
                if not line.startswith('#'):
                    first_paragraph.append(line)
                if len(first_paragraph) >= 2:
                    break
            
            if first_paragraph:
                summary = ' '.join(first_paragraph)
                if len(summary) > 300:
                    summary = summary[:297] + '...'
                return summary
        
        return None


class BaseParser:
    """Base class for all resource parsers"""
    
    def __init__(self):
        self.analyzer = ContentAnalyzer()
    
    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the file"""
        raise NotImplementedError
    
    def parse(self, file_path: str) -> ParsedResource:
        """Parse the file and return a ParsedResource"""
        raise NotImplementedError
    
    def _create_metadata(self, file_path: str, content: str) -> ResourceMetadata:
        """Create enhanced metadata for a file"""
        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        
        mime_type, _ = mimetypes.guess_type(file_path)
        
        created_at = None
        modified_at = None
        try:
            stat = os.stat(file_path)
            created_at = datetime.fromtimestamp(stat.st_ctime).isoformat()
            modified_at = datetime.fromtimestamp(stat.st_mtime).isoformat()
        except:
            pass
        
        quality, keywords, summary = self.analyzer.analyze_content(content, file_path)
        
        return ResourceMetadata(
            file_name=file_name,
            file_size=file_size,
            file_type=ext.lstrip('.'),
            mime_type=mime_type,
            created_at=created_at,
            modified_at=modified_at,
            content_length=len(content),
            content_quality=quality,
            keywords=keywords,
            summary=summary
        )


class TextParser(BaseParser):
    """Parser for plain text files"""
    
    TEXT_EXTENSIONS = {'.txt', '.md', '.markdown', '.rst', '.csv', '.json', '.yaml', '.yml', '.xml'}
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.TEXT_EXTENSIONS
    
    def parse(self, file_path: str) -> ParsedResource:
        ext = os.path.splitext(file_path)[1].lower()
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        resource_type = ResourceType.TEXT
        if ext in {'.md', '.markdown'}:
            resource_type = ResourceType.MARKDOWN
        elif ext == '.csv':
            resource_type = ResourceType.CSV
        elif ext in {'.json'}:
            resource_type = ResourceType.JSON
        elif ext in {'.yaml', '.yml'}:
            resource_type = ResourceType.YAML
        elif ext == '.xml':
            resource_type = ResourceType.XML
        
        metadata = self._create_metadata(file_path, content)
        
        return ParsedResource(
            file_path=file_path,
            resource_type=resource_type,
            text_content=content,
            metadata=metadata
        )


class PDFParser(BaseParser):
    """Parser for PDF files"""
    
    def can_parse(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def parse(self, file_path: str) -> ParsedResource:
        text_content = ""
        images = []
        custom_fields = {}
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                custom_fields['num_pages'] = len(reader.pages)
                if reader.metadata:
                    custom_fields['title'] = getattr(reader.metadata, 'title', '')
                    custom_fields['author'] = getattr(reader.metadata, 'author', '')
                
                for page in reader.pages:
                    text_content += page.extract_text() + "\n\n"
        
        except ImportError:
            text_content = "[PyPDF2 not installed. Install with: pip install PyPDF2]"
        except Exception as e:
            text_content = f"[Error parsing PDF: {str(e)}]"
        
        metadata = self._create_metadata(file_path, text_content)
        metadata.custom_fields = custom_fields
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.PDF,
            text_content=text_content,
            metadata=metadata,
            images=images
        )


class DOCParser(BaseParser):
    """Parser for DOC/DOCX files"""
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in {'.doc', '.docx'}
    
    def parse(self, file_path: str) -> ParsedResource:
        text_content = ""
        custom_fields = {}
        tables = []
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.docx':
                from docx import Document
                doc = Document(file_path)
                
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
                    text_content += "\n"
            
            else:
                text_content = "[.doc parsing requires additional setup. Consider converting to .docx first.]"
        
        except ImportError:
            text_content = "[python-docx not installed. Install with: pip install python-docx]"
        except Exception as e:
            text_content = f"[Error parsing document: {str(e)}]"
        
        metadata = self._create_metadata(file_path, text_content)
        metadata.custom_fields = custom_fields
        
        resource_type = ResourceType.DOCX if ext == '.docx' else ResourceType.DOC
        
        return ParsedResource(
            file_path=file_path,
            resource_type=resource_type,
            text_content=text_content,
            metadata=metadata,
            tables=tables
        )


class ExcelParser(BaseParser):
    """Parser for Excel files"""
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in {'.xlsx', '.xls', '.csv'}
    
    def parse(self, file_path: str) -> ParsedResource:
        text_content = ""
        custom_fields = {}
        tables = []
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            import pandas as pd
            
            if ext == '.csv':
                df = pd.read_csv(file_path)
                text_content = df.to_string(index=False)
                custom_fields['num_rows'] = len(df)
                custom_fields['num_columns'] = len(df.columns)
                custom_fields['columns'] = list(df.columns)
                
                tables.append([list(df.columns)] + df.values.tolist())
            
            else:
                xl = pd.ExcelFile(file_path)
                custom_fields['sheet_names'] = xl.sheet_names
                
                for sheet_name in xl.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    text_content += f"=== Sheet: {sheet_name} ===\n"
                    text_content += df.to_string(index=False)
                    text_content += "\n\n"
                    tables.append([list(df.columns)] + df.values.tolist())
        
        except ImportError:
            text_content = "[pandas not installed. Install with: pip install pandas openpyxl]"
        except Exception as e:
            text_content = f"[Error parsing Excel: {str(e)}]"
        
        metadata = self._create_metadata(file_path, text_content)
        metadata.custom_fields = custom_fields
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.EXCEL,
            text_content=text_content,
            metadata=metadata,
            tables=tables
        )


class ImageParser(BaseParser):
    """Parser for image files"""
    
    IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff'}
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.IMAGE_EXTENSIONS
    
    def parse(self, file_path: str) -> ParsedResource:
        custom_fields = {}
        
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                custom_fields['width'] = img.width
                custom_fields['height'] = img.height
                custom_fields['format'] = img.format
                custom_fields['mode'] = img.mode
        except ImportError:
            pass
        except Exception as e:
            custom_fields['error'] = str(e)
        
        text_content = f"Image file: {os.path.basename(file_path)}\n"
        text_content += f"Size: {custom_fields.get('width', 'unknown')}x{custom_fields.get('height', 'unknown')}"
        
        metadata = self._create_metadata(file_path, text_content)
        metadata.custom_fields = custom_fields
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.IMAGE,
            text_content=text_content,
            metadata=metadata,
            images=[raw_data],
            raw_data=raw_data
        )


class VideoParser(BaseParser):
    """Parser for video files"""
    
    VIDEO_EXTENSIONS = {'.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv'}
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.VIDEO_EXTENSIONS
    
    def parse(self, file_path: str) -> ParsedResource:
        custom_fields = {}
        
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        try:
            import cv2
            cap = cv2.VideoCapture(file_path)
            if cap.isOpened():
                custom_fields['fps'] = cap.get(cv2.CAP_PROP_FPS)
                custom_fields['frame_count'] = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                custom_fields['width'] = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                custom_fields['height'] = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                duration = custom_fields['frame_count'] / custom_fields['fps'] if custom_fields['fps'] > 0 else 0
                custom_fields['duration_seconds'] = duration
                cap.release()
        except ImportError:
            pass
        except Exception as e:
            custom_fields['error'] = str(e)
        
        text_content = f"Video file: {os.path.basename(file_path)}\n"
        text_content += f"Duration: {custom_fields.get('duration_seconds', 'unknown')} seconds\n"
        text_content += f"Resolution: {custom_fields.get('width', 'unknown')}x{custom_fields.get('height', 'unknown')}"
        
        metadata = self._create_metadata(file_path, text_content)
        metadata.custom_fields = custom_fields
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.VIDEO,
            text_content=text_content,
            metadata=metadata,
            raw_data=raw_data
        )


class HTMLParser(BaseParser):
    """Parser for HTML files"""
    
    def can_parse(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in {'.html', '.htm'}
    
    def parse(self, file_path: str) -> ParsedResource:
        custom_fields = {}
        links = []
        structured_sections = {}
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        text_content = html_content
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text_content = soup.get_text(separator='\n', strip=True)
            
            title = soup.title.string if soup.title else ''
            custom_fields['title'] = title
            
            for a in soup.find_all('a', href=True):
                link_text = a.get_text(strip=True) or a['href']
                links.append((link_text, a['href']))
            
            for heading in soup.find_all(['h1', 'h2', 'h3']):
                heading_text = heading.get_text(strip=True)
                if heading_text:
                    next_elem = heading.next_sibling
                    section_content = []
                    while next_elem and next_elem.name not in ['h1', 'h2', 'h3']:
                        if hasattr(next_elem, 'get_text'):
                            section_content.append(next_elem.get_text(strip=True))
                        next_elem = next_elem.next_sibling
                    if section_content:
                        structured_sections[heading_text] = '\n'.join(section_content)
        
        except ImportError:
            title = ''
        except Exception as e:
            text_content = f"[Error parsing HTML: {str(e)}]"
            title = ''
        
        metadata = self._create_metadata(file_path, text_content)
        metadata.custom_fields = custom_fields
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.HTML,
            text_content=text_content,
            metadata=metadata,
            structured_sections=structured_sections,
            links=links
        )


class ResourceParser:
    """Main resource parser that dispatches to appropriate parsers"""
    
    def __init__(self):
        self.parsers = [
            TextParser(),
            PDFParser(),
            DOCParser(),
            ExcelParser(),
            ImageParser(),
            VideoParser(),
            HTMLParser()
        ]
    
    def parse(self, file_path: str) -> ParsedResource:
        """
        Parse a file using the appropriate parser
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            ParsedResource object
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser.parse(file_path)
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        metadata = None
        try:
            temp_parser = TextParser()
            temp_parser.analyzer = ContentAnalyzer()
            metadata = temp_parser._create_metadata(file_path, content)
        except:
            file_size = os.path.getsize(file_path)
            metadata = ResourceMetadata(
                file_name=os.path.basename(file_path),
                file_size=file_size,
                file_type=os.path.splitext(file_path)[1].lstrip('.') or 'unknown',
                content_length=len(content),
                content_quality=ContentQuality.LOW
            )
        
        return ParsedResource(
            file_path=file_path,
            resource_type=ResourceType.UNKNOWN,
            text_content=content,
            metadata=metadata
        )
    
    def parse_to_skill_resource(self, file_path: str) -> str:
        """
        Parse a file and convert directly to skill resource format
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            String in skill resource format
        """
        parsed = self.parse(file_path)
        return parsed.to_skill_resource()


def parse_resource(file_path: str) -> ParsedResource:
    """
    Convenience function to parse a resource
    
    Args:
        file_path: Path to the file
        
    Returns:
        ParsedResource
    """
    parser = ResourceParser()
    return parser.parse(file_path)


def parse_to_skill_resource(file_path: str) -> str:
    """
    Convenience function to parse a file to skill resource format
    
    Args:
        file_path: Path to the file
        
    Returns:
        String in skill resource format
    """
    parser = ResourceParser()
    return parser.parse_to_skill_resource(file_path)
